import argparse
import csv
import json
import re
from pathlib import Path

from docx import Document


QUESTION_RE = re.compile(r"^\s*(\d{1,4})[.、．]?\s*(.*)")
SOURCE_RE = re.compile(r"^\s*\[(.+?)\]\s*$")
ANSWER_LINE_RE = re.compile(r"^\s*(作答|答案|标记|標記)[:：]")
INLINE_OPTION_RE = re.compile(r"(?:^|\s)[(（]?([A-Da-d])[)）.．、]\s*")


def read_docx_text(path):
    document = Document(path)
    lines = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" ".join(cells))

    return lines


def read_answers(path):
    if not path:
        return {}

    answers = {}
    with open(path, "r", encoding="utf-8-sig", newline="") as file:
        reader = csv.DictReader(file)
        for row in reader:
            question_id = str(row.get("id", "")).strip()
            answer = str(row.get("answer", "")).strip().upper()
            if question_id and answer:
                answers[question_id] = answer
    return answers


def split_inline_options(text):
    matches = list(INLINE_OPTION_RE.finditer(text))
    if not matches:
        return text.strip(), {}

    question_text = text[: matches[0].start()].strip()
    options = {}
    for index, match in enumerate(matches):
        letter = match.group(1).upper()
        start = match.end()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        options[letter] = text[start:end].strip()
    return question_text, options


def parse_option_line(line):
    match = re.match(r"^\s*[(（]?([A-Da-d])[)）.．、]\s*(.+?)\s*$", line)
    if not match:
        return None
    return match.group(1).upper(), match.group(2).strip()


def flush_question(question, questions, answers):
    if not question:
        return

    question["question"] = " ".join(part for part in question["question_parts"] if part).strip()
    question.pop("question_parts", None)
    question["answer"] = answers.get(str(question["id"]), question.get("answer", ""))

    if question["question"] and question["options"]:
        questions.append(question)


def parse_lines(lines, answers):
    questions = []
    current = None
    pending_source = ""

    for line in lines:
        if ANSWER_LINE_RE.match(line):
            continue

        source_match = SOURCE_RE.match(line)
        if source_match:
            pending_source = source_match.group(1).strip()
            continue

        question_match = QUESTION_RE.match(line)
        option = parse_option_line(line)

        if question_match and not option:
            flush_question(current, questions, answers)
            question_id = int(question_match.group(1))
            rest = question_match.group(2).strip()
            question_text, inline_options = split_inline_options(rest)
            current = {
                "id": question_id,
                "source": pending_source,
                "question_parts": [question_text],
                "options": inline_options,
                "answer": "",
            }
            pending_source = ""
            continue

        if not current:
            continue

        inline_question_text, inline_options = split_inline_options(line)
        if inline_options:
            if inline_question_text:
                current["question_parts"].append(inline_question_text)
            current["options"].update(inline_options)
            continue

        option = parse_option_line(line)
        if option:
            letter, text = option
            current["options"][letter] = text
            continue

        current["question_parts"].append(line.strip())

    flush_question(current, questions, answers)
    return questions


def write_json(questions, output_path):
    with open(output_path, "w", encoding="utf-8") as file:
        json.dump(questions, file, ensure_ascii=False, indent=2)


def main():
    parser = argparse.ArgumentParser(description="把 Word 选择题解析成 HTML 练习页可导入的 JSON。")
    parser.add_argument("docx", help="输入 Word 文件，例如 questions.docx")
    parser.add_argument("-a", "--answers", help="答案 CSV，表头为 id,answer")
    parser.add_argument("-o", "--output", default="questions.json", help="输出 JSON 文件")
    args = parser.parse_args()

    lines = read_docx_text(args.docx)
    answers = read_answers(args.answers)
    questions = parse_lines(lines, answers)
    write_json(questions, args.output)

    print(f"已生成 {args.output}，共 {len(questions)} 题。")
    missing_answers = [item["id"] for item in questions if not item.get("answer")]
    if missing_answers:
        print("以下题目没有匹配到答案：", ", ".join(map(str, missing_answers)))


if __name__ == "__main__":
    main()
